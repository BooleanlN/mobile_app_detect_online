# -*- coding:utf-8 -*-
"""
@author liuning
@date 2019/5/27 0:15
@File utils.py 
@Desciption 功能函数
"""
import os
import shutil
import tempfile

import time
import zipfile
from glob import glob

from docx import Document
import requests


ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}
ALLOWED_APK = {'apk','zip'}
IMAGE_RECO_URL = "http://192.168.139.139:5000/predict"
TEXT_CHECK_URL = "http://127.0.0.1:8777/filter"
LOCAL_DIRECTORY_PATH = os.getcwd() + "/" #当前文件路径
DOCS_SAVE_PATH = LOCAL_DIRECTORY_PATH + "result/docs/"
SOURCE_ZIP_PATH = LOCAL_DIRECTORY_PATH + "result/resoures/"

def unzip_apk_plus_identify(apk_temp_name):
    """
    apk解压加图片资源鉴别
    :param apk_temp_name: 
    :return: 
    """
    # 解压
    temp_apk_folder = unzip_file(apk_temp_name, "res")
    # 获取图片资源文件
    res = walk_apk_folders(temp_apk_folder)
    try:
        shutil.rmtree(temp_apk_folder)
    except  OSError:
        print(OSError.errno)
    os.remove(apk_temp_name)
    test = Document()
    test.add_heading('鉴别结果报告', 0)
    IMAGE_PATH = glob(res + "/*.*")
    s = requests.Session()
    s.keep_alive = False
    for IMAGE in sorted(IMAGE_PATH):
        test.add_heading("检测图片:" + IMAGE.split('/')[-1] + '\n', 2)
        image = open(IMAGE, "rb").read()
        payload = {"image": image}
        r = s.post(IMAGE_RECO_URL, files=payload)
        r = r.json()
        r_word = r["predictions"][0]
        # print(r)
        if r["success"]:
            test.add_paragraph("识别图片中的文字结果:" + r_word + "\n")
            text_reco_result = requests.post(TEXT_CHECK_URL, data={'content': r_word}).json()
            if(text_reco_result['success']):
                for word in text_reco_result['words']:
                    test.add_paragraph("敏感词：" + word+ "\n")
                test.add_paragraph("检测结果：" + text_reco_result["result"][0] + "\n")
    docsname = 'Result' + str(time.time()) + '.docx'
    test.save(DOCS_SAVE_PATH + docsname)
    return docsname
def is_img(filename):
    """
    :param filename: 
    :return: bool，校验图片格式
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS

def is_apk(filename):
    """
    :param filename: 
    :return: bool,校验apk格式
    """
    return '.' in filename and \
         filename.rsplit('.', 1)[1] in ALLOWED_APK

def get_image_content(filename):
  """
  :param filename: 
  :return: 图像识别接口
  """
  def get_file_content(path):
        # 二进制读取文件，读取完自动close
        with  open(path, 'rb') as fp:
            return fp.read()
  image = get_file_content(filename)
  os.remove(filename) # 删除
  payload = {"image":image}
  r = requests.post(url=IMAGE_RECO_URL,files=payload).json()
  if(r['success']):
    return r['predictions']
  else:
    return "图片格式错误"

def un_zip_apk(apk_temp_name):
    """
    :param apk_temp_name: 
    :param resname: 
    :return: 解压apk包
    """
    # 解压
    temp_apk_folder = unzip_file(apk_temp_name,"res")
    # 获取图片资源文件
    res = walk_apk_folders(temp_apk_folder)
    try:
        shutil.rmtree(temp_apk_folder)
    except  OSError:
        print(OSError.errno)
    images_ziped = add_dir_file(res) #打包文件
    os.remove(apk_temp_name)
    return images_ziped

def walk_apk_folders(file_dir):
    """
    :param file_dir: 
    :return:  图片文件夹路径，获取APK内的图片资源文件
    """
    IMG_TYPES=[".png",".jpg",".jpeg",".gif",".bmp"]
    IMG_DIR = "drawable"
    IMG_FOLDER = tempfile.mkdtemp("img","unzip",LOCAL_DIRECTORY_PATH+'temp/')
    #遍历文件夹
    for ass_root,ass_dir,ass_files in os.walk(file_dir+"/assets"):
      for file in ass_files:
        if (os.path.splitext(file)[1] in IMG_TYPES):
          try:
            shutil.move(file_dir + '/assets/' +file, IMG_FOLDER)
          except:
            continue
    for root,dirs,files in os.walk(file_dir+"/res"):
        #获得文件夹列表
        #print(files)
        for i in range(len(dirs)):
            #找到带有drawable的文件夹，疑似资源文件夹
            if dirs[i].find(IMG_DIR)!=-1:
                #遍历drawable文件夹
                for subroot, subdirs, subfiles in os.walk(file_dir+'/res/'+dirs[i]):
                    #print(subroot)
                    for j in range(len(subfiles)):
                        #获取到png、jpeg等文件
                        if (os.path.splitext(subfiles[j])[1] in IMG_TYPES):
                            try:
                                shutil.move(file_dir+'/res/'+dirs[i]+'/'+subfiles[j],IMG_FOLDER)
                            except:
                                continue
                        else:
                            os.remove(file_dir + '/res/' + dirs[i] + '/' + subfiles[j])
    return IMG_FOLDER

def add_dir_file(startdir):
    """
    :param startdir: 图片文件夹路径
    :return: 打包后的文件路径，将图片资源文件打包
    """

    timer = str(time.time())
    file_ = SOURCE_ZIP_PATH+timer+".zip"
    imgs_path_name = timer+".zip"
    f = zipfile.ZipFile(file_,'w',zipfile.ZIP_DEFLATED)
    for dirpath, dirnames, filenames in os.walk(startdir):
        for filename in filenames:
            f.write(os.path.join(dirpath,filename))
            os.remove(os.path.join(dirpath,filename))
    f.close()
    return imgs_path_name

def unzip_file(apk_temp_name,resname="res"):
    """
    解压缩apk
    :param apk_temp_name: apk文件路径
    :param resname: 结果文件夹
    :return: 结果文件夹路径
    """
    zip_file = zipfile.ZipFile(apk_temp_name)
    temp_apk_folder = tempfile.mkdtemp(resname + "_unzip", prefix="temp", dir=LOCAL_DIRECTORY_PATH + 'temp/apk/')
    for names in zip_file.namelist():
        try:
            zip_file.extract(names, temp_apk_folder)
        except:
            continue
    zip_file.close()
    return temp_apk_folder