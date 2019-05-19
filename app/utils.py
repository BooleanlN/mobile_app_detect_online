import os
import requests
# -*- coding: UTF-8 -*-
import time
import zipfile
import tempfile
import shutil
import uuid
from docx import Document
from glob import glob
PWD = '/var/www/html/'
def get_file_content(path):
  # 二进制读取文件，读取完自动close
  with  open(path, 'rb') as fp:
    return fp.read()
def get_image_content(filename):
  URL = "http://192.168.139.139:5000/predict"
  pwd = os.getcwd()
  res_dir = pwd + 'res/res' + str(time.time()) + '.txt'
  dir = 'res' + str(time.time()) + '.txt'
  image = get_file_content(filename)
  payload = {"image":image}
  r = requests.post(url=URL,files=payload).json()
  if(r['success']):
    return r['predictions']
  else:
    return "Image format error"
#只获取资源文件
def un_zip_unget(filename,resname="res"):
    zip_file = zipfile.ZipFile(filename)
    #pwd = os.getcwd()
    tempfolder=tempfile.mkdtemp(resname+"_unzip",prefix="temp",dir=PWD+'temp/apk/')
    print(tempfolder)
    for names in zip_file.namelist():
        try:
            zip_file.extract(names,tempfolder)
        except:
            continue
    zip_file.close()
    res = walkDir(tempfolder)
    try:
        shutil.rmtree(tempfolder)
    except  OSError:
        print(OSError.errno)
    file_ = adddirfile(res)
    os.remove(filename)
    return file_
#获取资源文件并识别
def un_zip(filename,resname="res"):
    zip_file = zipfile.ZipFile(filename)
    #pwd = os.getcwd()
    tempfolder=tempfile.mkdtemp(resname+"_unzip",prefix="temp",dir=PWD+'temp/apk/')
    #print(tempfolder)
    for names in zip_file.namelist():
        try:
            zip_file.extract(names,tempfolder)
        except:
            continue
    zip_file.close()
    res = walkDir(tempfolder)
    try:
        shutil.rmtree(tempfolder)
    except  OSError:
        print(OSError.errno)
    reco_txt = get_txt(res)
    file_ = adddirfile(res)
    os.remove(filename)
    return file_,reco_txt
#规范apk文件查找图像资源文件
def walkDir(file_dir):
    print(file_dir)
    Img=[".png",".jpg",".jpeg",".gif",".bmp"]
    ImgDir = "drawable"
    asset = "assets"
    ImgFolder = tempfile.mkdtemp("img","unzip",PWD+'temp/')
    #遍历文件夹
    for ass_root,ass_dir,ass_files in os.walk(file_dir+"/assets"):
      for file in ass_files:
        if (os.path.splitext(file)[1] in Img):
          try:
            shutil.move(file_dir + '/assets/' +file, ImgFolder)
          except:
            continue
    for root,dirs,files in os.walk(file_dir+"/res"):
        #获得文件夹列表
        #print(files)
        for i in range(len(dirs)):
            #找到带有drawable的文件夹，疑似资源文件夹
            if dirs[i].find(ImgDir)!=-1:
                #遍历drawable文件夹
                for subroot, subdirs, subfiles in os.walk(file_dir+'/res/'+dirs[i]):
                    #print(subroot)
                    for j in range(len(subfiles)):
                        #获取到png、jpeg等文件
                        if (os.path.splitext(subfiles[j])[1] in Img):
                            try:
                                shutil.move(file_dir+'/res/'+dirs[i]+'/'+subfiles[j],ImgFolder)
                            except:
                                continue
    return ImgFolder
#把整个文件夹内的文件打包
def adddirfile(startdir):
    pwd = PWD+"res/ImageZip/"
    timer = str(time.time())
    file_ = pwd+timer+".zip"
    return_name = "ImageZip/"+timer+".zip"
    f = zipfile.ZipFile(file_,'w',zipfile.ZIP_DEFLATED)
    for dirpath, dirnames, filenames in os.walk(startdir):
        for filename in filenames:
            f.write(os.path.join(dirpath,filename))
            os.remove(os.path.join(dirpath,filename))
    f.close()
    return return_name
def text_detect(str):
  r = requests.post(TEXT_URL,{'content': str}).json()
  print(r)
def getWord(filename,resname="res"):
  zip_file = zipfile.ZipFile(filename)
  #pwd = os.getcwd()
  tempfolder = tempfile.mkdtemp(resname + "_unzip", prefix="temp", dir=PWD + 'temp/apk/')
  #print(tempfolder)
  for names in zip_file.namelist():
    try:
      zip_file.extract(names, tempfolder)
    except:
      continue
  zip_file.close()
  res = walkDir(tempfolder)
  try:
    shutil.rmtree(tempfolder)
  except  OSError:
    print(OSError.errno)
  os.remove(filename)
  test = Document()
  test.add_heading('鉴别结果报告', 0)
  IMAGE_PATH = glob(res+"/*.*")
  #图片识别接口
  KERAS_REST_API_URL = "http://192.168.139.139:5000/predict"
  #KERAS_API_URL = "http://54.186.163.83:5000/filter"
  FILE_SAVE_PATH = PWD+"res/wordRes/"
  s = requests.Session()
  s.keep_alive = False
  for IMAGE in sorted(IMAGE_PATH):
    test.add_heading("检测图片:" + IMAGE.split('/')[-1] + '\n', 2)
    image = open(IMAGE, "rb").read()
    payload = {"image": image}
    r = s.post(KERAS_REST_API_URL, files=payload)
    r = r.json()
    r_word = ""
    # ensure the request was sucessful
    if r["success"]:
      #loop over the predictions and display them
      test.add_paragraph("识别图片中的文字结果:" + r["reco_words"]+"\n")
      test.add_paragraph("敏感词："+r["words"]+"\n")
      test.add_paragraph("检测结果："+r["result"])
    # load the input image and construct the payload for the request
  timer = str(time.time())
  test.save(FILE_SAVE_PATH+'Result'+timer+'.docx')
  return "wordRes/"+'Result'+timer+'.docx'
def getWords(filename,resname="res"):
  zip_file = zipfile.ZipFile(filename)
  #pwd = os.getcwd()
  tempfolder = tempfile.mkdtemp(resname + "_unzip", prefix="temp", dir=PWD + 'temp/apk/')
  #print(tempfolder)
  for names in zip_file.namelist():
    try:
      zip_file.extract(names, tempfolder)
    except:
      continue
  zip_file.close()
  res = walkDir(tempfolder)
  try:
    shutil.rmtree(tempfolder)
  except  OSError:
    print(OSError.errno)
  os.remove(filename)
  test = Document()
  test.add_heading('鉴别结果报告', 0)
  IMAGE_PATH = glob(res+"/*.*")
  KERAS_REST_API_URL = "http://52.24.169.228:5000/predict"
  KERAS_API_URL = "http://54.186.163.83:5000/filter"
  FILE_SAVE_PATH = PWD+"res/wordRes/"
  s = requests.Session()
  s.keep_alive = False
  for IMAGE in sorted(IMAGE_PATH):
    test.add_heading("检测图片:" + IMAGE.split('/')[-1] + '\n', 2)
    image = open(IMAGE, "rb").read()
    payload = {"image": image}
    r = s.post(KERAS_REST_API_URL, files=payload)
    r = r.json()
    r_word = ""
    # ensure the request was sucessful
    if r["success"]:
      # loop over the predictions and display them
      for i in r["predictions"]:
        r_word += i
    test.add_paragraph("识别图片中的文字结果:" + r_word)
    words = ""
    # load the input image and construct the payload for the request
    payload = {'content': r_word}
    # submit the request
    r = s.post(KERAS_API_URL, data=payload).json()
    if r["success"]:
      # loop over the predictions and display them
      for i in r["words"]:
        words += i
      words += ","
      if words == "":
        test.add_paragraph("敏感词:无")
      else:
        test.add_paragraph("敏感词:" + words)
      for j in r["result"]:
        test.add_paragraph("结果鉴定:" + j + '\n')
  timer = str(time.time())
  test.save(FILE_SAVE_PATH+'Result'+timer+'.docx')
  return "wordRes/"+'Result'+timer+'.docx'
if __name__ == '__main__':
  pwd = os.getcwd()
  print("请求开始...")
  #get_image_content(pwd + '/4.png')
  #text_detect('xxx')
  getWord(os.getcwd()+'/temp/unzip0yn0ltyaimg')
